"""
Rule-based ATS-compatibility checker for uploaded CVs.

Runs entirely in memory: the uploaded file's bytes are parsed for text and
structure inside analyze(), and nothing is ever written to disk or the
database. Once the request finishes, Python garbage-collects the bytes —
no CV content or personal data persists anywhere on the server.
"""
from __future__ import annotations
import io
import re
from collections import Counter
from dataclasses import dataclass

MAX_FILE_BYTES = 3 * 1024 * 1024
MAX_PDF_PAGES = 8  # a real resume is 1-3 pages; this bounds worst-case memory on oversized uploads
ALLOWED_EXTENSIONS = {"pdf", "docx"}


class UnsupportedFile(Exception):
    pass


@dataclass
class Check:
    id: str
    title: str
    passed: bool
    severity: str  # "critical" | "warning" | "tip"
    detail: str
    fix: str
    points: int
    max_points: int


def _ext(filename: str) -> str:
    return (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()


def _extract_pdf(data: bytes) -> tuple[str, bool, bool]:
    import pdfplumber

    text_parts: list[str] = []
    has_tables = False
    has_images = False
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages[:MAX_PDF_PAGES]:
            text_parts.append(page.extract_text() or "")
            if page.extract_tables():
                has_tables = True
            if page.images:
                has_images = True
            # pdfplumber caches each page's decoded layout/image objects on
            # the Page object as it goes; on a memory-constrained VM this
            # adds up fast across pages, so release it as soon as we're done
            # reading this page rather than waiting for the `with` to exit.
            page.flush_cache()
    return "\n".join(text_parts), has_tables, has_images


def _extract_docx(data: bytes) -> tuple[str, bool, bool]:
    from docx import Document

    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    return text, len(doc.tables) > 0, len(doc.inline_shapes) > 0


EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_RE = re.compile(r"(\+?\d[\d\s().-]{7,}\d)")
BULLET_RE = re.compile(r"(^|\n)\s*[•\-\*▪◦‣]\s+", re.M)
SECTION_PATTERNS = {
    "experience": re.compile(r"\b(work experience|professional experience|experience|employment history)\b", re.I),
    "education": re.compile(r"\b(education|academic background)\b", re.I),
    "skills": re.compile(r"\b(skills|technical skills|core competencies)\b", re.I),
}


def _run_checks(text: str, has_tables: bool, has_images: bool) -> list[Check]:
    checks: list[Check] = []
    word_count = len(text.split())

    parseable = word_count >= 30
    checks.append(Check(
        "parseable", "Text is machine-readable", parseable,
        "critical",
        "Your CV's text extracted cleanly." if parseable else "We could barely extract any text from this file.",
        "This file may be a scanned image or built entirely from graphics/text boxes. Rebuild it as "
        "a real text document (export from Word, or a text-based PDF) — ATS software can't read pictures of text.",
        25 if parseable else 0, 25,
    ))

    has_email = bool(EMAIL_RE.search(text))
    checks.append(Check(
        "email", "Email address found", has_email,
        "critical" if not has_email else "tip",
        "Found a valid-looking email address." if has_email else "No email address detected.",
        "Add a clear email address near the top of your CV, in plain text — not inside an image, header, or text box.",
        10 if has_email else 0, 10,
    ))

    has_phone = bool(PHONE_RE.search(text))
    checks.append(Check(
        "phone", "Phone number found", has_phone,
        "warning" if not has_phone else "tip",
        "Found a phone number." if has_phone else "No phone number detected.",
        "Add a phone number next to your other contact details so recruiters can reach you directly.",
        5 if has_phone else 0, 5,
    ))

    for key, pts in (("experience", 10), ("education", 5), ("skills", 5)):
        found = bool(SECTION_PATTERNS[key].search(text))
        checks.append(Check(
            f"section_{key}", f"'{key.title()}' section present", found,
            "warning" if not found else "tip",
            f"Found a clearly labeled {key} section." if found else f"We couldn't find a clearly labeled {key} section.",
            f"Add a section titled '{key.title()}' — ATS software parses resumes by these standard "
            "headings, and creative alternatives (like \"What I've Built\") often get missed entirely.",
            pts if found else 0, pts,
        ))

    checks.append(Check(
        "tables", "No tables used for layout", not has_tables,
        "warning",
        "No tables detected." if not has_tables else "This CV uses one or more tables.",
        "Replace tables with plain paragraphs and line breaks. Many ATS parsers read table cells out "
        "of order or skip them entirely, scrambling your work history.",
        15 if not has_tables else 0, 15,
    ))

    checks.append(Check(
        "images", "No embedded images or graphics", not has_images,
        "warning",
        "No embedded images detected." if not has_images else "This CV contains embedded images or graphics.",
        "Remove logos, icons, and graphical skill bars. ATS software ignores images entirely, and any "
        "text inside them (like a photo of your contact details) is invisible to it.",
        10 if not has_images else 0, 10,
    ))

    length_ok = 250 <= word_count <= 1400
    if length_ok:
        length_detail = f"{word_count} words — a good length."
    elif word_count < 250:
        length_detail = f"{word_count} words — too short to show much depth."
    else:
        length_detail = f"{word_count} words — quite long; ATS and recruiters both skim."
    checks.append(Check(
        "length", "Reasonable length", length_ok,
        "tip", length_detail,
        "Aim for roughly one to two pages (about 400-900 words) — enough to show depth without losing the reader.",
        10 if length_ok else 0, 10,
    ))

    has_bullets = bool(BULLET_RE.search(text))
    checks.append(Check(
        "bullets", "Uses bullet points", has_bullets,
        "tip",
        "Bullet points found." if has_bullets else "No bullet points detected.",
        "Use bullet points for your responsibilities and achievements — dense paragraphs are harder "
        "for both ATS software and human reviewers to scan quickly.",
        5 if has_bullets else 0, 5,
    ))

    return checks


STOPWORDS = {
    "a", "about", "above", "after", "again", "against", "all", "am", "an", "and", "any", "are", "aren't",
    "as", "at", "be", "because", "been", "before", "being", "below", "between", "both", "but", "by",
    "can", "cannot", "could", "did", "do", "does", "doing", "down", "during", "each", "few", "for",
    "from", "further", "had", "has", "have", "having", "he", "her", "here", "hers", "herself", "him",
    "himself", "his", "how", "i", "if", "in", "into", "is", "it", "its", "itself", "let's", "me", "more",
    "most", "my", "myself", "nor", "of", "on", "once", "only", "or", "other", "ought", "our", "ours",
    "ourselves", "out", "over", "own", "same", "she", "should", "so", "some", "such", "than", "that",
    "the", "their", "theirs", "them", "themselves", "then", "there", "these", "they", "this", "those",
    "through", "to", "too", "under", "until", "up", "very", "was", "we", "were", "what", "when", "where",
    "which", "while", "who", "whom", "why", "with", "would", "you", "your", "yours", "yourself",
    "yourselves", "will", "shall", "also", "etc", "using", "use", "used", "including", "include", "per",
    "role", "job", "work", "team", "years", "year", "strong", "ability", "experience", "required",
}
WORD_RE = re.compile(r"[a-zA-Z][a-zA-Z+#.\-]{1,}")


def _significant_words(text: str) -> list[str]:
    words = [w.lower().strip(".-") for w in WORD_RE.findall(text)]
    return [w for w in words if w and w not in STOPWORDS and len(w) > 2]


def _keyword_match(resume_text: str, job_description: str) -> dict | None:
    jd_words = _significant_words(job_description)
    if not jd_words:
        return None
    resume_words = set(_significant_words(resume_text))
    top_keywords = [w for w, _ in Counter(jd_words).most_common(40)]
    matched = [w for w in top_keywords if w in resume_words]
    missing = [w for w in top_keywords if w not in resume_words][:12]
    score = round(len(matched) / len(top_keywords) * 100) if top_keywords else 0
    return {"score": score, "matched": matched[:20], "missing": missing}


def analyze(filename: str, data: bytes, job_description: str = "") -> dict:
    if len(data) > MAX_FILE_BYTES:
        raise UnsupportedFile("That file is larger than 3MB — please upload a smaller one.")
    ext = _ext(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise UnsupportedFile("Please upload a .pdf or .docx file.")

    try:
        if ext == "pdf":
            text, has_tables, has_images = _extract_pdf(data)
        else:
            text, has_tables, has_images = _extract_docx(data)
    except Exception as e:
        raise UnsupportedFile(
            "Couldn't read this file — it may be corrupted or password-protected."
        ) from e

    checks = _run_checks(text, has_tables, has_images)
    score = sum(c.points for c in checks)
    max_score = sum(c.max_points for c in checks)

    result = {
        "score": round(score / max_score * 100) if max_score else 0,
        "word_count": len(text.split()),
        "checks": [
            {
                "id": c.id, "title": c.title, "passed": c.passed, "severity": c.severity,
                "detail": c.detail, "fix": c.fix,
            }
            for c in checks
        ],
    }
    if job_description.strip():
        result["keyword_match"] = _keyword_match(text, job_description)
    return result
